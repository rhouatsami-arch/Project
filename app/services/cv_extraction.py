"""
app/services/cv_extraction.py
-------------------------------
Production-grade CV text extraction pipeline.

Pipeline: file bytes -> validate -> extract raw text -> parse structured data
Supports: PDF (pdfplumber), DOCX (python-docx)
"""
import re
import json
import magic
from typing import Optional


ALLOWED_MIME_TYPES = {
    "application/pdf":    "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
}

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

SKILL_KEYWORDS = [
    "python", "java", "javascript", "typescript", "react", "vue", "angular",
    "fastapi", "django", "flask", "node", "express", "sql", "postgresql",
    "mysql", "mongodb", "docker", "kubernetes", "git", "aws", "azure", "gcp",
    "machine learning", "deep learning", "nlp", "ai", "tensorflow", "pytorch",
    "html", "css", "sass", "rest api", "graphql", "ci/cd", "agile", "scrum",
    "linux", "bash", "powershell", "excel", "tableau", "power bi",
]


def detect_real_filetype(file_bytes: bytes) -> Optional[str]:
    """Detect actual file type from content (not just extension) — security check."""
    mime = magic.from_buffer(file_bytes, mime=True)
    return ALLOWED_MIME_TYPES.get(mime)


def validate_cv_upload(file_bytes: bytes, filename: str) -> tuple[bool, str]:
    """Returns (is_valid, error_message)."""
    if len(file_bytes) > MAX_FILE_SIZE:
        return False, "File exceeds 5MB limit"
    if len(file_bytes) == 0:
        return False, "Empty file"

    real_ext = detect_real_filetype(file_bytes)
    if not real_ext:
        return False, "Unsupported or corrupted file type — only PDF/DOC/DOCX allowed"

    claimed_ext = filename.rsplit(".", 1)[-1].lower()
    if claimed_ext not in ("pdf", "doc", "docx"):
        return False, "Invalid file extension"

    return True, ""


def extract_text_from_pdf(file_path: str) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # also extract text from tables (common in CV templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_raw_text(file_path: str, ext: str) -> str:
    """Dispatch extraction by file type. Never raises — returns empty string on failure."""
    try:
        if ext == "pdf":
            return extract_text_from_pdf(file_path)
        elif ext in ("doc", "docx"):
            return extract_text_from_docx(file_path)
    except Exception:
        return ""
    return ""


def parse_structured_data(text: str) -> dict:
    """Extract structured signals from raw CV text."""
    emails = list(set(re.findall(r"[\w.+-]+@[\w-]+\.\w+", text)))
    phones = list(set(re.findall(r"(?:\+\d{1,3}[\s.-]?)?\(?\d{2,4}\)?[\s.-]?\d{2,4}[\s.-]?\d{2,4}[\s.-]?\d{0,4}", text)))
    phones = [p.strip() for p in phones if len(re.sub(r"\D", "", p)) >= 8][:5]
    urls   = list(set(re.findall(r"https?://[^\s,;]+", text)))

    text_lower = text.lower()
    skills_detected = sorted({kw for kw in SKILL_KEYWORDS if kw in text_lower})

    return {
        "emails":           emails[:3],
        "phones":           phones[:3],
        "urls":              urls[:5],
        "skills_detected":  skills_detected,
        "word_count":       len(text.split()),
        "raw_text_preview": text[:800],
    }


def process_cv(file_path: str, ext: str) -> dict:
    """Full pipeline: extract + parse. Returns dict ready for DB storage."""
    raw_text = extract_raw_text(file_path, ext)
    if not raw_text.strip():
        return {"success": False, "raw_text": "", "extracted_data": {}, "error": "No text could be extracted"}

    structured = parse_structured_data(raw_text)
    return {
        "success":        True,
        "raw_text":       raw_text,
        "extracted_data": structured,
        "error":          None,
    }
